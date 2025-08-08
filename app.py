
import streamlit as st
import re
from PyPDF2 import PdfFileReader
from PIL import Image
import pytesseract
from pdf2image import convert_from_bytes
from collections import defaultdict
import docx
from io import BytesIO
from docx import Document
from fpdf import FPDF

# Palavras-chave
PALAVRAS_CHAVE = [
    "Esclarecimento", "Impugnação", "data de abertura", "recursos",
    "prazo", "horas", "garantia", "caução", "seguro-garantia",
    "entrega", "atestado de capacidade técnica", "regional", "local"
]

def encontrar_contexto(texto, palavra, n_caracteres=120):
    trechos = []
    for match in re.finditer(re.escape(palavra), texto, flags=re.IGNORECASE):
        start = max(match.start() - n_caracteres, 0)
        end = min(match.end() + n_caracteres, len(texto))
        contexto = texto[start:end].replace('\n', ' ').strip()
        trechos.append(contexto)
    return trechos

st.title("🔍 Analisador de Palavras-chave em Documentos")
st.markdown("Faça upload de arquivos PDF, Word ou imagens e clique em **Gerar Relatório**.")

uploaded_files = st.file_uploader("📁 Upload de documentos", type=["pdf", "jpg", "jpeg", "docx"], accept_multiple_files=True)

relatorios = []

if st.button("🚀 Gerar Relatório") and uploaded_files:
    with st.spinner("🔄 Processando arquivos..."):
        for file in uploaded_files:
            st.markdown("---")
            texto = ''
            try:
                ext = file.name.lower().split(".")[-1]
                if ext == "pdf":
                    leitor = PdfFileReader(file)
                    for i in range(leitor.numPages):
                        pagina = leitor.getPage(i)
                        texto_pagina = pagina.extractText()
                        if texto_pagina:
                            texto += texto_pagina
                    if not texto.strip():
                        imagens = convert_from_bytes(file.read())
                        for img in imagens:
                            texto += pytesseract.image_to_string(img, lang='por+eng') + '\n'

                elif ext in ["jpg", "jpeg"]:
                    imagem = Image.open(file)
                    texto = pytesseract.image_to_string(imagem, lang='por+eng')

                elif ext == "docx":
                    doc = docx.Document(file)
                    texto = '\n'.join([p.text for p in doc.paragraphs])

                texto_limpo = texto.strip()
                resumo = texto_limpo[:500] + '...' if texto_limpo else '(Sem texto extraído)'

                st.subheader(f"📄 Arquivo: {file.name}")
                st.markdown("🔍 **Prévia do Texto:**")
                st.code(resumo)

                relatorio = [f"📄 Arquivo: {file.name}\n", f"🔍 Prévia do Texto:\n{resumo}\n"]

                for palavra in PALAVRAS_CHAVE:
                    trechos = encontrar_contexto(texto_limpo, palavra)
                    if trechos:
                        st.markdown(f"➡️ **Palavra:** `{palavra}` | Ocorrências: `{len(trechos)}`")
                        relatorio.append(f"➡️ Palavra: '{palavra}' | Ocorrências: {len(trechos)}\n")
                        for i, trecho in enumerate(trechos[:3], 1):
                            st.markdown(f"- ✏️ **Contexto {i}:** ...{trecho}...")
                            relatorio.append(f"   ✏️ Contexto {i}: ...{trecho}...\n")
                relatorios.append('\n'.join(relatorio))
            except Exception as e:
                st.error(f"Erro ao processar {file.name}: {e}")

   if relatorios:
    opcao = st.selectbox("📤 Exportar como:", ["Word", "PDF"])
    
    if st.button("💾 Gerar Arquivo"):
        nome_arquivo = "relatorio_final"
        
        if opcao == "Word":
            doc = Document()
            for rel in relatorios:
                doc.add_paragraph(rel)
                doc.add_paragraph("------------------------------------------------")
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            st.download_button("📥 Baixar Word", data=buffer, file_name=f"{nome_arquivo}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        elif opcao == "PDF":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for rel in relatorios:
                for line in rel.splitlines():
                    pdf.multi_cell(0, 10, line)
                pdf.cell(0, 10, "-" * 60, ln=True)
            pdf_bytes = pdf.output(dest='S').encode('latin-1')
            buffer = BytesIO(pdf_bytes)
            st.download_button("📥 Baixar PDF", data=buffer, file_name=f"{nome_arquivo}.pdf", mime="application/pdf")


        elif opcao == "PDF":
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            for rel in relatorios:
                for line in rel.splitlines():
                    pdf.multi_cell(0, 10, line)
                pdf.cell(0, 10, "-" * 60, ln=True)
            # Corrigido: usar output(dest='S').encode('latin-1')
            pdf_bytes = pdf.output(dest='S').encode('latin-1')
            buffer = BytesIO(pdf_bytes)
            st.download_button("📥 Baixar PDF", data=buffer, file_name=f"{nome_arquivo}.pdf", mime="application/pdf")
